from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, CallbackQueryHandler
import datetime
import os

# Mengatur path ke file database (db.txt)
db_path = './db.txt'

# Load the document
doc_path = './penawaran.docx'
doc = Document(doc_path)

# Fungsi untuk mendapatkan nomor bulan Romawi
def get_roman_month(month):
    roman_months = {
        1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI',
        7: 'VII', 8: 'VIII', 9: 'IX', 10: 'X', 11: 'XI', 12: 'XII'
    }
    return roman_months.get(month, '')

# Fungsi untuk mendapatkan nomor form yang berikutnya
def get_next_form_number(db_path, current_year):
    # Mengecek apakah file database (db.txt) ada atau tidak
    # if not os.path.exists(db_path) or not os.path.isfile(db_path):
    #     # Jika tidak ada, mulai dari nomor 1
    #     return 1
    
    # Membaca file untuk mendapatkan nomor terakhir yang digunakan
    with open(db_path, 'r') as file:
        last_entry = file.read().strip()
    
    # Memeriksa apakah file kosong atau tahun telah berubah
    if not last_entry or (2000 + int(last_entry.split('/')[-1])) != current_year:
        print(last_entry)
        print(current_year)
        # Jika tahun telah berubah atau file kosong, mulai dari nomor 1
        return 1
    
    # Jika tahun sama, lanjutkan ke nomor selanjutnya
    last_number = int(last_entry.split('/')[0])
    return last_number + 1

# Fungsi untuk menyimpan nomor form yang baru ke database (db.txt)
def save_form_number(db_path, form_number, year):
    with open(db_path, 'w') as file:
        file.write(f"{str(form_number).zfill(2)}/FR/BM/{get_roman_month(datetime.datetime.now().month)}/{str(year)[2:]}")
    

def generateEditedDocument():
    # Assume we need to add a dynamic number to the "No:" field
    # For demonstration, let's use a hypothetical dynamic number
    # Mendapatkan tanggal saat ini
    current_date = datetime.datetime.now()
    current_year = current_date.year

    # Mendapatkan nomor form berikutnya
    next_form_number = get_next_form_number(db_path, current_year)

    # Menyimpan nomor form yang baru
    save_form_number(db_path, next_form_number, current_year)

    # Membuat format nomor form
    dynamic_number = f"{str(next_form_number).zfill(2)}/FR/BM/{get_roman_month(current_date.month)}/{str(current_year)[2:]}"
    print(dynamic_number)

    # Search for the paragraph containing "No:" and add the dynamic number
    for paragraph in doc.paragraphs:
        if "No:" in paragraph.text:
            # Clear the existing run(s) after "No:" in the paragraph
            # Here we assume "No:" is followed by a space and then the dynamic content
            for run in paragraph.runs:
                if "No:" in run.text:
                    text_split = run.text.split('No:')
                    run.text = text_split[0] + 'No: ' + dynamic_number
                    break

    # Save the edited document
    edited_doc_path = './bm-form-registration-training-2024.docx'
    doc.save(edited_doc_path)

# edited_doc_path

def send_document(update, context):
    query = update.callback_query
    chat_id = query.message.chat_id
    generateEditedDocument()
    context.bot.send_document(chat_id, document=open('./bm-form-registration-training-2024.docx', 'rb'))

# Token API dari BotFather
# token = '6815627895:AAGU_kQ6A3w3lKT7P_nf5R-rlRGLq7UEyAc'

# updater = Updater(token)
# dp = updater.dispatcher

# # Menambahkan command handler untuk mengirim dokumen
# dp.add_handler(CommandHandler('cetak_registrasi', send_document))

# updater.start_polling()
# updater.idle()

# Fungsi yang dipanggil ketika pengguna menekan tombol start
def start(update, context):
    keyboard = [
        [InlineKeyboardButton("Cetak Registrasi", callback_data='cetak_registrasi')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    update.message.reply_text('Halo Warga Brainmatics! Kamu mau ngapain nih?? Klik Tombol dibawah yaa :)', reply_markup=reply_markup)

# Fungsi yang dipanggil ketika pengguna menekan tombol Cetak Registrasi
def button(update, context):
    query = update.callback_query
    query.answer()
    
    # Memeriksa data callback dan memanggil fungsi yang sesuai
    if query.data == 'cetak_registrasi':
        send_document(update, context)

# Fungsi utama untuk mengatur bot
def main():
    # Token API dari BotFather
    token = '6815627895:AAGU_kQ6A3w3lKT7P_nf5R-rlRGLq7UEyAc'
    updater = Updater(token)
    
    # Dispatcher untuk mendaftarkan handlers
    dp = updater.dispatcher
    
    # Menambahkan command handler untuk /start
    dp.add_handler(CommandHandler('start', start))
    
    # Menambahkan callback query handler untuk menangani tombol
    dp.add_handler(CallbackQueryHandler(button))
    
    # Mulai bot
    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()