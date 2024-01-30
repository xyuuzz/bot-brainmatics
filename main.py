from docx import Document
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, CallbackQueryHandler, MessageHandler, Filters
import datetime
import os
import pandas as pd
from send_email import SendEmail
from spreadsheet import Spreadsheet


# Mengatur path ke file database (db.txt)
db_path = './db.txt'

# Load the document
doc_path = './penawaran.docx'
doc = Document(doc_path)

link_google_keamanan = 'https://myaccount.google.com/u/1/security?hl=in'

# Token API dari BotFather
token = '6815627895:AAEIbmtmC4ByHgIrXyxODKo3J0MBgk8MInw'
updater = Updater(token)

# Dispatcher untuk mendaftarkan handlers
dp = updater.dispatcher

spr = Spreadsheet("https://docs.google.com/spreadsheets/d/1sG12Jf2wWQKkxx6TPkUuwUtwUY0xvilhgpRCxKo_oLg/edit?pli=1#gid=136109342", 'Nomor Form Registrasi 2024')

# Fungsi untuk mendapatkan nomor bulan Romawi
def get_roman_month(month):
    roman_months = {
        1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI',
        7: 'VII', 8: 'VIII', 9: 'IX', 10: 'X', 11: 'XI', 12: 'XII'
    }
    return roman_months.get(month, '')

# Fungsi untuk mendapatkan nomor form yang berikutnya
def get_next_form_number(last_entry, current_year):
    # Memeriksa apakah file kosong atau tahun telah berubah
    if not last_entry or (2000 + int(last_entry.split('/')[-1])) != current_year:
        # Jika tahun telah berubah atau file kosong, mulai dari nomor 1
        return 1
    
    # Jika tahun sama, lanjutkan ke nomor selanjutnya
    last_number = int(last_entry.split('/')[0])
    return last_number + 1

# Fungsi untuk menyimpan nomor form yang baru ke database (db.txt)
def save_form_number(db_path, form_number, year):
    with open(db_path, 'w') as file:
        file.write(f"{str(form_number).zfill(2)}/FR/BM/{get_roman_month(datetime.datetime.now().month)}/{str(year)[2:]}")
    

def generateEditedDocument(last_entry):
    # Assume we need to add a dynamic number to the "No:" field
    # For demonstration, let's use a hypothetical dynamic number
    # Mendapatkan tanggal saat ini
    # current_date = datetime.datetime.now()
    # current_year = current_date.year

    # Mendapatkan nomor form berikutnya
    # next_form_number = get_next_form_number(last_entry, current_year)

    # Menyimpan nomor form yang baru
    # save_form_number(db_path, next_form_number, current_year)

    # Membuat format nomor form
    dynamic_number = last_entry

    # Search for the paragraph containing "No:" and add the dynamic number
    for paragraph in doc.paragraphs:
        if "No:" in paragraph.text:
            # Clear the existing run(s) after "No:" in the paragraph
            # Here we assume "No:" is followed by a space and then the dynamic content
            for run in paragraph.runs:
                if "No:" in run.text:
                    text_split = run.text.split('No:')
                    run.text = text_split[0] + 'No: ' + dynamic_number
                    break

    # Save the edited document
    edited_doc_path = './bm-form-registration-training-2024.docx'

    doc.save(edited_doc_path)


def generateNoFR(latest_number):
    # Mendapatkan tanggal saat ini
    current_date = datetime.datetime.now()
    current_year = current_date.year

    # Mendapatkan nomor form berikutnya
    next_form_number = get_next_form_number(latest_number, current_year)

    # Membuat format nomor form
    dynamic_number = f"{str(next_form_number).zfill(2)}/FR/BM/{get_roman_month(current_date.month)}/{str(current_year)[2:]}"
    return dynamic_number

def send_document(update, context):
    query = update.callback_query
    chat_id = query.message.chat_id
    generateEditedDocument()
    context.bot.send_document(chat_id, document=open('./bm-form-registration-training-2024.docx', 'rb'))


# Fungsi yang dipanggil ketika pengguna menekan tombol start
def start(update, context):
    keyboard = [
        [
            InlineKeyboardButton("Dapatkan Nomor Form Registrasi", callback_data='cetak_registrasi'),
            InlineKeyboardButton("Kirim Email Konfirmasi Peserta Training", callback_data='kirim_email_konfirmasi')
        ],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    update.message.reply_text('Halo Warga Brainmatics! Kamu mau ngapain nih?? Klik Tombol dibawah yaa :)', reply_markup=reply_markup)

statusPesan = 'mulai'

# Fungsi yang dipanggil pertama kali
def button(update, context):
    global statusPesan

    query = update.callback_query
    chat_id = query.message.chat_id

    # cek jawaban user
    value = update.callback_query.data


    if(statusPesan == 'fr_pilih_jenis_training'):
        latest_number = spr.get_value_last_row()
        noFRBaru = generateNoFR(latest_number)
        spr.add_data([noFRBaru, value])

        # context.bot.send_message(chat_id, f'Baik, anda memilih jenis training *{value}*')
        context.bot.send_message(chat_id, 'Silahkan cantumkan Nama Training')
        statusPesan = 'fr_nama_training'
    
    # Memeriksa data callback dan memanggil fungsi yang sesuai
    if query.data == 'cetak_registrasi':
        selectJenisTraining(query, context)
    elif query.data == 'kirim_email_konfirmasi':
        statusPesan = 'mulai'
        context.bot.send_message(chat_id, 'Silahkan lengkapi data training dan peserta training pada file excel di bawah ini untuk mengirimkan Email Konfirmasi Training ke Peserta')
        context.bot.send_document(chat_id, document=open('./Template Data Email Konfirmasi.xlsx', 'rb'))
        context.bot.send_message(chat_id, 'Setelah melengkapi data training dan peserta training pada file excel, kirim kesini lagi ya filenya :)')
        context.bot.send_message(chat_id, 'Nanti BOT ini akan membantu mengirimkan secara otomatis Email Konfirmasi Training ke peserta training ya')
        context.bot.send_message(chat_id, 'NB: Aktifkan pengaturan Less Secure App Access (Akses Aplikasi Kurang Aman) pada akun email yang digunakan untuk mengirimkan Email Konfirmasi Training')
        context.bot.send_message(chat_id, f'Untuk mengaturnya, silahkan bisa klik link berikut ini: {link_google_keamanan}')
        context.bot.send_message(chat_id, 'Berikut untuk foto menu pengaturannya:')
        context.bot.send_photo(chat_id, photo=open('./ss_full_akses_aplikasi_kurang_aman.png', 'rb'))
        context.bot.send_photo(chat_id, photo=open('./akses_aplikasi_kurang_aman.png', 'rb'))






# Fungsi untuk menangani pemilihan
def selectJenisTraining(update, context):
    global statusPesan
    keyboard = [
        [InlineKeyboardButton("Regular Training", callback_data="Regular Training")],
        [InlineKeyboardButton("Private Training", callback_data="Private Training")],
        [InlineKeyboardButton("InHouse Training", callback_data="InHouse Training")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    statusPesan = 'fr_pilih_jenis_training'
    update.message.reply_text("Pilih jenis training:", reply_markup=reply_markup)

def handle_document(update, context):
    global statusPesan

    chat_id = update.message.chat_id
    if(statusPesan == 'mulai'):
        context.bot.send_message(chat_id, 'Terima kasih sudah mengirimkan data berupa file excel :>')
        context.bot.send_message(chat_id, 'Untuk langkah selanjutnya silahkan bisa mengirimkan file surat konfirmasi training berupa PDF')
        context.bot.send_message(chat_id, 'Berikut untuk template surat konfirmasi training yang bisa digunakan')
        context.bot.send_document(chat_id, document=open('template surat konfirmasi training.docx', 'rb'))

        statusPesan = 'upload_dokumen_training'

        # Mendapatkan file_id dari dokumen yang dikirim
        file_id = update.message.document.file_id
        new_file = context.bot.get_file(file_id)

        # Mengunduh file
        file_path = new_file.file_path
        new_file.download('file_data.xlsx')
        return
    elif(statusPesan == 'upload_dokumen_training'):
        # Mendapatkan file_id dari dokumen yang dikirim
        file_id = update.message.document.file_id
        new_file = context.bot.get_file(file_id)

        # Mengunduh file
        file_path = new_file.file_path
        new_file.download('surat_konfirmasi.pdf')

    context.bot.send_message(chat_id, 'Loading.., wait a minute :>')

    # Membaca file Excel
    df = pd.read_excel('file_data.xlsx')

    # Proses data di sini, misalnya membaca data dari kolom tertentu
    # index ke 0 adalah baris dan index ke 1 adalah kolom
    # data_kolom = df.iloc[0, 0]

    # excel = MappingExcel(df)
    # excel.mappingData()

    namaTraining = df.iloc[0, 1]
    tanggalTraining = df.iloc[1, 1]
    waktuTraining = df.iloc[2, 1]
    lokasiTraining = df.iloc[3, 1]
    linkLokasiGmaps = df.iloc[4, 1]
    ruanganTraining = df.iloc[5, 1]
    jumlahPeserta = df.iloc[6, 1]

    namaAsisten = df.iloc[10, 1]
    noHpAsisten = df.iloc[12, 1]

    emailAsisten = df.iloc[11, 1]
    passwordAkunBrainmaticsAsisten = df.iloc[14, 1]

    # password akun menggunakan app password (nanti ada tutorial cara membuatnya)
    sd = SendEmail(emailAsisten, passwordAkunBrainmaticsAsisten)
    sd.setSmtpSettings('smtp.gmail.com', 587)

    file_path = sd.getAttachmentPath(namaTraining)

    ccEmails = ['info@brainmatics.com']
    # ccEmails = []

    totalSoftwarePerluDisiapkan = df.iloc[1, 8]
    linkDownloadSoftware = df.iloc[2, 8]
    listSoftware = []

    if(totalSoftwarePerluDisiapkan > 0):
        for i in range(7, totalSoftwarePerluDisiapkan+7):
            listSoftware.append(df.iloc[i, 7])


    # looping array banyak data peserta
    # mulai dari angka 2 karena data peserta mulai dari baris ke 3
    for i in range(2, jumlahPeserta+2):
        namaPeserta = df.iloc[i, 2]
        emailPeserta = df.iloc[i, 3]
        usernamePeserta = df.iloc[i, 4]
        passwordPeserta = df.iloc[i, 5]

        bodyEmail = sd.getBodyEmail(namaAsisten, noHpAsisten, namaPeserta, namaTraining,
                                    tanggalTraining, waktuTraining, lokasiTraining,
                                    linkLokasiGmaps, ruanganTraining, usernamePeserta, passwordPeserta,
                                    totalSoftwarePerluDisiapkan, linkDownloadSoftware, listSoftware)

        sd.send(f"Konfirmasi Pelaksanaan {namaTraining}", bodyEmail, emailPeserta, ccEmails, namaAsisten, file_path)
        context.bot.send_message(chat_id, f"Email Konfirmasi Training berhasil dikirimkan ke *{namaPeserta}* melalui email *{emailPeserta}*")

    context.bot.send_message(chat_id, 'Semua email konfirmasi peserta berhasil dikirim!')

    os.remove('file_data.xlsx')
    os.remove('surat_konfirmasi.pdf')
    os.remove(file_path)

def handle_text(update, context):
    global statusPesan
    chat_id = update.message.chat_id
    last_row = spr.get_last_row()

    # menerima value
    value = update.message.text

    if (statusPesan.split('_').pop(0) == 'fr'):
        column = '_'.join(statusPesan.split('_')[1:])
        spr.update_data('registrasi', column, last_row, value)

    if(statusPesan == 'fr_nama_training'):
        context.bot.send_message(chat_id, 'Silahkan cantumkan Instansi yang ingin melakukan Registrasi Training')
        statusPesan = 'fr_instansi'
    elif(statusPesan == 'fr_instansi'):
        context.bot.send_message(chat_id, 'Silahkan cantumkan nama PIC Internal (Marketing)')
        statusPesan = 'fr_pic_internal'
    elif(statusPesan == 'fr_pic_internal'):
        context.bot.send_message(chat_id, 'Silahkan cantumkan nama PIC Eksternal (PIC Instansi/Client)')
        statusPesan = 'fr_pic_eksternal'
    elif (statusPesan == 'fr_pic_eksternal'):
        # context.bot.send_message(chat_id, 'Terimakasih atas data yang diberikan! Berikut untuk Nomor Form Registrasi Trainingnya ya staff marketing BM tercintaah :>')
        #context.bot.send_message(chat_id, spr.get_value_last_row())
        context.bot.send_message(chat_id, 'Terimakasih atas data yang diberikan! Berikut untuk Dokumen Form Registrasi Trainingnya ya staff marketing BM tercintaah :>')
        last_entry = spr.get_value_last_row()
        generateEditedDocument(last_entry)
        context.bot.send_document(chat_id, document=open('./bm-form-registration-training-2024.docx', 'rb'))
        context.bot.send_message(chat_id, 'Semoga training nya jadi ya, biar bisa dapet komisi :>')
        statusPesan = 'mulai'


# Fungsi utama untuk mengatur bot
def main():
    global statusPesan
    # Menambahkan command handler untuk /start
    dp.add_handler(CommandHandler('start', start))
    
    # Menambahkan callback query handler untuk menangani tombol
    dp.add_handler(CallbackQueryHandler(button))

    # Menambahkan message handler untuk menangani dokumen/file
    dp.add_handler(MessageHandler(Filters.document, handle_document))

    # Menambahkan message handler untuk menangani text
    dp.add_handler(MessageHandler(Filters.text, handle_text))

    # Mulai bot
    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()